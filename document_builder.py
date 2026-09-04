"""
document_builder.py
Formats the output state produced by the LangGraph pipeline (state_graph.py)
into a professional Legal Brief Word document, and exports the same data
as a structured JSON fact sheet.
"""

from __future__ import annotations

import json
import os
from datetime import datetime
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.environ.get("LEGAL_OUTPUT_DIR", "output")

_ACCENT_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
_MUTED_COLOR = RGBColor(0x5A, 0x5A, 0x5A)
_HIGH_SEVERITY_COLOR = RGBColor(0xB3, 0x00, 0x1B)
_MED_SEVERITY_COLOR = RGBColor(0xB3, 0x7A, 0x00)
_LOW_SEVERITY_COLOR = RGBColor(0x2E, 0x7D, 0x32)


def _ensure_output_dir() -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    return OUTPUT_DIR


def _add_title_page(document: Document, state: Dict[str, Any]) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LEGAL BRIEF & COURT SUBMISSION PACKAGE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = _ACCENT_COLOR

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    jurisdiction = state.get("jurisdiction", {})
    sub_run = subtitle.add_run(
        f"{jurisdiction.get('court_type', 'General')} Matter — "
        f"{jurisdiction.get('court_level', 'Court Level Pending')}"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = _MUTED_COLOR

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Source Document: {state.get('file_name', 'N/A')}    |    "
        f"Client Representation: {state.get('client_role', 'N/A')}    |    "
        f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = _MUTED_COLOR

    document.add_paragraph()
    _add_horizontal_rule(document)
    document.add_page_break()


def _add_horizontal_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("─" * 90)
    run.font.color.rgb = _ACCENT_COLOR
    run.font.size = Pt(8)


def _add_section_heading(document: Document, number: int, title: str) -> None:
    heading = document.add_heading(level=1)
    run = heading.add_run(f"{number}. {title}")
    run.font.color.rgb = _ACCENT_COLOR
    run.font.size = Pt(16)


def _add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text or "No data available.")
    paragraph.paragraph_format.space_after = Pt(10)
    for run in paragraph.runs:
        run.font.size = Pt(11)


def _severity_color(severity: str) -> RGBColor:
    normalized = (severity or "").strip().lower()
    if normalized == "high":
        return _HIGH_SEVERITY_COLOR
    if normalized == "medium":
        return _MED_SEVERITY_COLOR
    return _LOW_SEVERITY_COLOR


def _section_1_executive_summary(document: Document, state: Dict[str, Any]) -> None:
    _add_section_heading(document, 1, "Executive Summary & Fact Sheet")
    _add_body_paragraph(document, state.get("executive_summary", ""))

    document.add_heading("Parties Identified", level=2)
    parties = state.get("parties", [])
    if parties:
        table = document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_cells = table.rows[0].cells
        for idx, heading_text in enumerate(["Name", "Role", "Description"]):
            header_cells[idx].text = heading_text
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for party in parties:
            row_cells = table.add_row().cells
            row_cells[0].text = str(party.get("name", ""))
            row_cells[1].text = str(party.get("role", ""))
            row_cells[2].text = str(party.get("description", ""))
    else:
        _add_body_paragraph(document, "No parties were identified in the source document.")
    document.add_paragraph()


def _section_2_party_positions(document: Document, state: Dict[str, Any]) -> None:
    _add_section_heading(document, 2, "Party Positions & Dispute Analysis")
    _add_body_paragraph(document, state.get("party_positions", ""))


def _section_3_timeline_and_audit(document: Document, state: Dict[str, Any]) -> None:
    _add_section_heading(document, 3, "Chronological Event Timeline & Lie/Contradiction Audit")

    document.add_heading("Event Timeline", level=2)
    timeline = state.get("timeline", [])
    if timeline:
        table = document.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        for idx, heading_text in enumerate(["Date", "Event", "Source Snippet", "Disputed"]):
            header_cells[idx].text = heading_text
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for event in timeline:
            row_cells = table.add_row().cells
            row_cells[0].text = str(event.get("date", ""))
            row_cells[1].text = str(event.get("event", ""))
            row_cells[2].text = str(event.get("source_snippet", ""))[:300]
            row_cells[3].text = "Yes" if event.get("disputed") else "No"
    else:
        _add_body_paragraph(document, "No dated events were extracted from the source document.")

    document.add_paragraph()
    document.add_heading("Inconsistency / Contradiction Audit", level=2)
    inconsistencies = state.get("inconsistencies", [])
    if inconsistencies:
        for item in inconsistencies:
            bullet = document.add_paragraph(style="List Bullet")
            severity = item.get("severity", "Low")
            severity_run = bullet.add_run(f"[{severity.upper()} SEVERITY] ")
            severity_run.bold = True
            severity_run.font.color.rgb = _severity_color(severity)
            bullet.add_run(item.get("description", ""))
            for statement in item.get("conflicting_statements", []):
                sub_bullet = document.add_paragraph(style="List Bullet 2")
                sub_run = sub_bullet.add_run(f"\u201c{statement}\u201d")
                sub_run.italic = True
    else:
        _add_body_paragraph(document, "No inconsistencies or contradictions were flagged.")
    document.add_paragraph()


def _section_4_jurisdiction(document: Document, state: Dict[str, Any]) -> None:
    _add_section_heading(document, 4, "Jurisdictional Routing & Court Fee Schedule")
    jurisdiction = state.get("jurisdiction", {})

    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    rows_data = [
        ("Court Level", jurisdiction.get("court_level", "N/A")),
        ("Court Type", jurisdiction.get("court_type", "N/A")),
        ("Estimated Court Fee", jurisdiction.get("estimated_court_fee", "N/A")),
    ]
    for label, value in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        row_cells[1].text = str(value)

    document.add_paragraph()
    document.add_heading("Routing Rationale", level=2)
    _add_body_paragraph(document, jurisdiction.get("reasoning", ""))

    document.add_heading("Filing Requirements", level=2)
    requirements = jurisdiction.get("filing_requirements", [])
    if requirements:
        for requirement in requirements:
            document.add_paragraph(requirement, style="List Bullet")
    else:
        _add_body_paragraph(document, "No specific filing requirements were determined.")
    document.add_paragraph()


def _section_5_precedents(document: Document, state: Dict[str, Any]) -> None:
    _add_section_heading(document, 5, "Cited Precedents & Legal Authorities")
    precedents = state.get("precedents", [])
    if precedents:
        for precedent in precedents:
            paragraph = document.add_paragraph()
            name_run = paragraph.add_run(f"{precedent.get('case_name', 'Unnamed Case')} ")
            name_run.bold = True
            citation_run = paragraph.add_run(f"[{precedent.get('citation', 'No citation')}]")
            citation_run.italic = True
            citation_run.font.color.rgb = _MUTED_COLOR

            support_paragraph = document.add_paragraph()
            support_label = "Supports Client Position" if precedent.get("supports_client") else "Adverse / Distinguishable"
            support_run = support_paragraph.add_run(f"{support_label}: ")
            support_run.bold = True
            support_paragraph.add_run(precedent.get("relevance", ""))
            document.add_paragraph()
    else:
        _add_body_paragraph(document, "No precedents were matched for this case.")


def _section_6_formal_speech(document: Document, state: Dict[str, Any]) -> None:
    _add_section_heading(document, 6, "Formal Court Oral Argument / Speech")
    speech_text = state.get("formal_speech", "")
    if speech_text:
        for paragraph_text in speech_text.split("\n"):
            if paragraph_text.strip():
                paragraph = document.add_paragraph(paragraph_text.strip())
                paragraph.paragraph_format.space_after = Pt(10)
                paragraph.paragraph_format.line_spacing = 1.5
    else:
        _add_body_paragraph(document, "No formal speech was generated.")


def build_legal_brief_document(state: Dict[str, Any]) -> Document:
    """Builds a python-docx Document object from the pipeline's final state."""
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    _add_title_page(document, state)
    _section_1_executive_summary(document, state)
    document.add_page_break()
    _section_2_party_positions(document, state)
    document.add_page_break()
    _section_3_timeline_and_audit(document, state)
    document.add_page_break()
    _section_4_jurisdiction(document, state)
    document.add_page_break()
    _section_5_precedents(document, state)
    document.add_page_break()
    _section_6_formal_speech(document, state)

    return document


def build_fact_sheet_json(state: Dict[str, Any]) -> Dict[str, Any]:
    """Builds the structured JSON fact sheet from the pipeline's final state."""
    return {
        "file_id": state.get("file_id"),
        "file_name": state.get("file_name"),
        "client_role": state.get("client_role"),
        "generated_at": datetime.now().isoformat(),
        "executive_summary": state.get("executive_summary", ""),
        "party_positions": state.get("party_positions", ""),
        "parties": state.get("parties", []),
        "timeline": state.get("timeline", []),
        "inconsistencies": state.get("inconsistencies", []),
        "jurisdiction": state.get("jurisdiction", {}),
        "precedents": state.get("precedents", []),
        "formal_speech": state.get("formal_speech", ""),
        "processing_errors": state.get("errors", []),
        "status": state.get("status", "unknown"),
    }


def save_case_outputs(state: Dict[str, Any]) -> Dict[str, str]:
    """
    Builds and saves both the .docx legal brief and the .json fact sheet
    for a given pipeline state, keyed by state['file_id'].

    Returns a dict with the absolute paths of both saved files.
    """
    output_dir = _ensure_output_dir()
    file_id = state.get("file_id") or datetime.now().strftime("%Y%m%d%H%M%S")

    docx_path = os.path.join(output_dir, f"{file_id}_legal_brief.docx")
    json_path = os.path.join(output_dir, f"{file_id}_fact_sheet.json")

    document = build_legal_brief_document(state)
    document.save(docx_path)

    fact_sheet = build_fact_sheet_json(state)
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(fact_sheet, json_file, indent=2, ensure_ascii=False)

    return {"docx_path": docx_path, "json_path": json_path}
