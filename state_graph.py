"""
state_graph.py
LangGraph architecture for the Legal Case Analysis & Court Speech Generation System.

Defines a StateGraph with four sequential nodes:
    1. DocumentParserNode
    2. TimelineAndPartyExtractorNode
    3. JurisdictionAndPrecedentNode
    4. LegalDraftingNode
"""
from __future__ import annotations
from dotenv import load_dotenv
load_dotenv()

import json
import os
import re
import struct
from datetime import datetime
from typing import Any, Dict, List, Optional, TypedDict

import fitz  # PyMuPDF
import docx as docx_lib
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from langgraph.graph import StateGraph, END
from pydantic import BaseModel, Field


# --------------------------------------------------------------------------
# LLM CONFIGURATION
# --------------------------------------------------------------------------

def get_llm(temperature: float = 0.1) -> ChatOpenAI:
    """Returns a configured ChatOpenAI client. Reads OPENAI_API_KEY from env."""
    model_name = os.environ.get("LEGAL_LLM_MODEL", "gpt-4o")
    return ChatOpenAI(
        model=model_name,
        temperature=temperature,
        api_key=os.environ.get("OPENAI_API_KEY"),
        max_retries=3,
        timeout=120,
    )


# --------------------------------------------------------------------------
# SHARED STATE SCHEMA
# --------------------------------------------------------------------------

class Party(TypedDict):
    name: str
    role: str  # "Plaintiff" | "Defendant" | "Witness" | "Third Party"
    description: str


class TimelineEvent(TypedDict):
    date: str
    event: str
    source_snippet: str
    disputed: bool


class Inconsistency(TypedDict):
    description: str
    conflicting_statements: List[str]
    severity: str  # "Low" | "Medium" | "High"


class Precedent(TypedDict):
    case_name: str
    citation: str
    relevance: str
    supports_client: bool


class JurisdictionInfo(TypedDict):
    court_level: str
    court_type: str
    reasoning: str
    estimated_court_fee: str
    filing_requirements: List[str]


class CaseState(TypedDict, total=False):
    # Ingestion
    file_path: str
    file_name: str
    raw_text: str
    client_role: str  # "Plaintiff / Victim" | "Defendant / Respondent"

    # TimelineAndPartyExtractorNode outputs
    parties: List[Party]
    timeline: List[TimelineEvent]
    inconsistencies: List[Inconsistency]

    # JurisdictionAndPrecedentNode outputs
    jurisdiction: JurisdictionInfo
    precedents: List[Precedent]

    # LegalDraftingNode outputs
    executive_summary: str
    party_positions: str
    formal_speech: str

    # Bookkeeping
    file_id: str
    errors: List[str]
    status: str


# --------------------------------------------------------------------------
# PYDANTIC SCHEMAS FOR STRUCTURED LLM OUTPUT
# --------------------------------------------------------------------------

class _PartyModel(BaseModel):
    name: str
    role: str = Field(description="Plaintiff, Defendant, Witness, or Third Party")
    description: str


class _TimelineEventModel(BaseModel):
    date: str = Field(description="Best-guess date in YYYY-MM-DD, or textual date if unclear")
    event: str
    source_snippet: str
    disputed: bool = False


class _InconsistencyModel(BaseModel):
    description: str
    conflicting_statements: List[str]
    severity: str = Field(description="Low, Medium, or High")


class _TimelineExtractionResult(BaseModel):
    parties: List[_PartyModel]
    timeline: List[_TimelineEventModel]
    inconsistencies: List[_InconsistencyModel]


class _PrecedentModel(BaseModel):
    case_name: str
    citation: str
    relevance: str
    supports_client: bool


class _JurisdictionModel(BaseModel):
    court_level: str = Field(description="e.g. High Court, District/Lower Court, Tribunal")
    court_type: str = Field(description="e.g. Civil, Commercial, Labor, Family, Criminal")
    reasoning: str
    estimated_court_fee: str
    filing_requirements: List[str]


class _JurisdictionExtractionResult(BaseModel):
    jurisdiction: _JurisdictionModel
    precedents: List[_PrecedentModel]


class _DraftingResult(BaseModel):
    executive_summary: str
    party_positions: str
    formal_speech: str


# --------------------------------------------------------------------------
# NODE 1: DocumentParserNode
# --------------------------------------------------------------------------

def DocumentParserNode(state: CaseState) -> CaseState:
    """
    Extracts structured plain text from an uploaded PDF, DOCX, DOC, or TXT file.
    Populates state["raw_text"].
    """
    errors = state.get("errors", [])
    file_path = state.get("file_path")
    file_name = state.get("file_name", "")

    if not file_path or not os.path.exists(file_path):
        errors.append(f"DocumentParserNode: file not found at '{file_path}'")
        return {**state, "raw_text": "", "errors": errors, "status": "parse_failed"}

    extension = os.path.splitext(file_name or file_path)[1].lower()
    extracted_text = ""

    try:
        if extension == ".pdf":
            extracted_text = _extract_pdf_text(file_path)
        elif extension == ".docx":
            extracted_text = _extract_docx_text(file_path)
        elif extension == ".doc":
            extracted_text = _extract_doc_text(file_path)
        elif extension == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                extracted_text = f.read()
        else:
            errors.append(f"DocumentParserNode: unsupported file extension '{extension}'")
    except Exception as exc:  # noqa: BLE001
        errors.append(f"DocumentParserNode: extraction error - {exc}")

    extracted_text = _normalize_whitespace(extracted_text)

    return {
        **state,
        "raw_text": extracted_text,
        "errors": errors,
        "status": "parsed" if extracted_text else "parse_failed",
    }


def _extract_pdf_text(file_path: str) -> str:
    text_chunks: List[str] = []
    with fitz.open(file_path) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            page_text = page.get_text("text")
            if page_text.strip():
                text_chunks.append(f"[Page {page_index}]\n{page_text.strip()}")
    return "\n\n".join(text_chunks)


def _extract_docx_text(file_path: str) -> str:
    document = docx_lib.Document(file_path)
    parts: List[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_doc_text(file_path: str) -> str:
    """
    Extracts plain text from legacy .doc (Microsoft Word 97-2003) binary files.
    
    Attempts structured parsing using olefile first. If that fails,
    falls back to extracting readable ASCII/UTF-8 text chunks from the binary stream.
    Never raises an exception; always returns a string (empty if parsing fails completely).
    
    Args:
        file_path: Path to the .doc file.
    
    Returns:
        Extracted plain text as a single string. Returns empty string if extraction fails.
    """
    # Attempt 1: Structured parsing via olefile
    try:
        import olefile
        
        if not olefile.isOleFile(file_path):
            # Not a valid OLE file; try binary fallback
            return _extract_doc_binary_fallback(file_path)
        
        ole = olefile.OleFileIO(file_path)
        try:
            # The main document content is in the "WordDocument" stream
            if ole.exists("WordDocument"):
                # Try to extract structured text from OLE streams
                doc_text = _parse_ole_word_document(ole)
                if doc_text.strip():
                    return doc_text
        except Exception:
            # Log but don't crash; try fallback
            pass
        finally:
            ole.close()
    except ImportError:
        # olefile not installed; skip to fallback
        pass
    except Exception:
        # Any error during OLE parsing; try fallback
        pass
    
    # Attempt 2: Binary fallback – extract readable text chunks
    return _extract_doc_binary_fallback(file_path)


def _parse_ole_word_document(ole) -> str:
    """
    Extracts text from a Word 97-2003 .doc file's OLE streams.
    Handles the complex .doc format by reading the WordDocument stream
    and extracting text from the main document body.
    
    Args:
        ole: An olefile.OleFileIO instance.
    
    Returns:
        Extracted text string.
    """
    try:
        # Read the WordDocument stream (main document metadata & pointers)
        word_doc_data = ole.openstream("WordDocument").read()
        
        if len(word_doc_data) < 0x1A:
            return ""
        
        # Determine if it's 0Table or 1Table based on FIB flags at offset 0x08
        flags = struct.unpack("<H", word_doc_data[0x08:0x0A])[0]
        table_name = "0Table" if (flags & 0x0200) else "1Table"
        
        text_parts = []
        
        # Try extracting readable text directly from streams
        for stream_name in ["WordDocument", table_name]:
            if ole.exists(stream_name):
                try:
                    stream_data = ole.openstream(stream_name).read()
                    readable = _extract_readable_text_from_stream(stream_data)
                    if readable:
                        text_parts.append(readable)
                except Exception:
                    pass
        
        if text_parts:
            return "\n".join(text_parts)
    except Exception:
        pass
    
    return ""


def _extract_doc_binary_fallback(file_path: str) -> str:
    """
    Fallback: extract readable ASCII and UTF-8 text chunks from the binary stream.
    This works for most .doc files even if structured parsing fails.
    
    Looks for sequences of printable ASCII and valid UTF-8 characters,
    filters out control codes, and joins coherent text runs.
    
    Args:
        file_path: Path to the .doc file.
    
    Returns:
        Extracted plain text as a string.
    """
    try:
        with open(file_path, "rb") as f:
            binary_data = f.read()
    except Exception:
        return ""
    
    if not binary_data:
        return ""
    
    text_chunks = []
    current_chunk = bytearray()
    
    i = 0
    while i < len(binary_data):
        byte = binary_data[i]
        
        # Printable ASCII range: 0x20–0x7E, plus common whitespace
        if 0x20 <= byte <= 0x7E or byte in (0x09, 0x0A, 0x0D):  # tab, newline, CR
            current_chunk.append(byte)
            i += 1
        # UTF-8 multi-byte sequences (0xC0–0xFD)
        elif 0xC0 <= byte <= 0xFD:
            # Attempt to decode a UTF-8 sequence
            utf8_seq = bytearray([byte])
            j = i + 1
            max_continuation = 3 if byte < 0xE0 else (2 if byte < 0xF0 else 1)
            
            while j < len(binary_data) and len(utf8_seq) <= max_continuation:
                next_byte = binary_data[j]
                # Continuation bytes: 10xxxxxx (0x80–0xBF)
                if 0x80 <= next_byte <= 0xBF:
                    utf8_seq.append(next_byte)
                    j += 1
                else:
                    break
            
            # Try to decode; if valid, include it
            try:
                decoded = utf8_seq.decode("utf-8")
                # Ensure no control characters (except whitespace)
                if not any(0 <= ord(c) < 0x20 and c not in "\t\n\r" for c in decoded):
                    current_chunk.extend(utf8_seq)
                i = j
            except (UnicodeDecodeError, AttributeError):
                # Not valid UTF-8; flush current chunk and skip this byte
                if current_chunk:
                    try:
                        chunk_text = current_chunk.decode("ascii", errors="ignore").strip()
                        if chunk_text and len(chunk_text) > 2:
                            text_chunks.append(chunk_text)
                    except Exception:
                        pass
                    current_chunk = bytearray()
                i += 1
        else:
            # Non-text byte; flush current chunk if non-empty
            if current_chunk:
                try:
                    chunk_text = current_chunk.decode("ascii", errors="ignore").strip()
                    if chunk_text and len(chunk_text) > 2:
                        text_chunks.append(chunk_text)
                except Exception:
                    pass
                current_chunk = bytearray()
            i += 1
    
    # Flush final chunk
    if current_chunk:
        try:
            chunk_text = current_chunk.decode("ascii", errors="ignore").strip()
            if chunk_text and len(chunk_text) > 2:
                text_chunks.append(chunk_text)
        except Exception:
            pass
    
    # Join chunks, collapse extra whitespace
    full_text = "\n".join(text_chunks)
    full_text = re.sub(r"\n{3,}", "\n\n", full_text)  # Collapse multiple newlines
    full_text = re.sub(r"[ \t]+", " ", full_text)      # Collapse multiple spaces/tabs
    
    return full_text.strip()


def _extract_readable_text_from_stream(data: bytes) -> str:
    """
    Simple helper to extract readable ASCII sequences from binary data.
    Used by the OLE parser when it needs to pull text directly from streams.
    
    Args:
        data: Binary stream data.
    
    Returns:
        Extracted text string.
    """
    # Find all sequences of printable ASCII (0x20–0x7E) + common whitespace
    text_runs = re.findall(rb'[\x20-\x7E\t\n\r]{4,}', data)
    decoded_runs = []
    
    for run in text_runs:
        try:
            decoded = run.decode("ascii", errors="ignore").strip()
            if decoded and len(decoded) > 2:  # Filter out very short fragments
                decoded_runs.append(decoded)
        except Exception:
            pass
    
    return "\n".join(decoded_runs)


def _normalize_whitespace(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --------------------------------------------------------------------------
# NODE 2: TimelineAndPartyExtractorNode
# --------------------------------------------------------------------------

def TimelineAndPartyExtractorNode(state: CaseState) -> CaseState:
    """
    Identifies parties (Plaintiff vs. Defendant), builds a chronological
    timeline of events, and flags inconsistencies/contradictions in the
    source document.
    """
    errors = state.get("errors", [])
    raw_text = state.get("raw_text", "")
    client_role = state.get("client_role", "Plaintiff / Victim")

    if not raw_text:
        errors.append("TimelineAndPartyExtractorNode: no raw_text available to analyze")
        return {**state, "parties": [], "timeline": [], "inconsistencies": [], "errors": errors}

    system_prompt = (
        "You are a meticulous paralegal AI. You extract structured facts from raw "
        "case documents (witness statements, complaints, contracts, correspondence). "
        "You never invent facts that are not supported by the text. "
        "The user's client is representing the '{client_role}' side of this dispute. "
        "Return only information that can be grounded in the provided text."
    ).format(client_role=client_role)

    user_prompt = (
        "Analyze the following case document text and extract:\n"
        "1. All parties involved, with their role (Plaintiff, Defendant, Witness, or "
        "Third Party) and a one-sentence description of their involvement.\n"
        "2. A chronological timeline of every dated or datable event mentioned, each "
        "with the exact source snippet it was drawn from. If a precise date is not "
        "given, provide the best textual approximation (e.g. 'early March 2024').\n"
        "3. Any inconsistencies or contradictions between statements in the document "
        "(e.g. conflicting dates, conflicting accounts of the same event), each rated "
        "Low/Medium/High severity, with the conflicting statements quoted.\n\n"
        f"DOCUMENT TEXT:\n{raw_text[:60000]}"
    )

    try:
        llm = get_llm(temperature=0.0)
        structured_llm = llm.with_structured_output(_TimelineExtractionResult)
        result: _TimelineExtractionResult = structured_llm.invoke(
            [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]
        )
        parties = [p.model_dump() for p in result.parties]
        timeline = sorted(
            [t.model_dump() for t in result.timeline],
            key=lambda item: _sortable_date(item.get("date", "")),
        )
        inconsistencies = [i.model_dump() for i in result.inconsistencies]
    except Exception as exc:  # noqa: BLE001
        errors.append(f"TimelineAndPartyExtractorNode: LLM extraction error - {exc}")
        parties, timeline, inconsistencies = [], [], []

    return {
        **state,
        "parties": parties,
        "timeline": timeline,
        "inconsistencies": inconsistencies,
        "errors": errors,
        "status": "timeline_extracted",
    }


def _sortable_date(date_str: str) -> str:
    """Best-effort normalization so timeline entries sort chronologically."""
    formats = ["%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%B %d, %Y", "%d %B %Y", "%Y-%m", "%Y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str.strip(), fmt).isoformat()
        except (ValueError, AttributeError):
            continue
    return date_str or "9999-99-99"


# --------------------------------------------------------------------------
# NODE 3: JurisdictionAndPrecedentNode
# --------------------------------------------------------------------------

_COURT_FEE_SCHEDULE = {
    "small_claims": "1% of claim value (min. $50)",
    "district_civil": "2% of claim value (min. $150)",
    "high_court_civil": "3% of claim value (min. $500)",
    "commercial": "2.5% of claim value (min. $750)",
    "labor_tribunal": "Flat fee: $75",
    "family_court": "Flat fee: $200",
}


def JurisdictionAndPrecedentNode(state: CaseState) -> CaseState:
    """
    Determines the appropriate court level and type, estimates applicable
    court fees, and matches relevant legal precedents with citations.
    """
    errors = state.get("errors", [])
    raw_text = state.get("raw_text", "")
    parties = state.get("parties", [])
    timeline = state.get("timeline", [])
    client_role = state.get("client_role", "Plaintiff / Victim")

    if not raw_text:
        errors.append("JurisdictionAndPrecedentNode: no raw_text available to analyze")
        return {**state, "jurisdiction": _empty_jurisdiction(), "precedents": [], "errors": errors}

    system_prompt = (
        "You are a legal research AI specializing in jurisdictional routing and case "
        "law research. Given case facts, you determine the correct court level and "
        "case type, estimate a plausible court filing fee, and identify legal "
        "precedents that are genuinely relevant to the fact pattern. For every "
        "precedent you cite, provide a case name and citation in standard legal "
        "citation format, and explain in one sentence why it is relevant and whether "
        "it supports or undermines the client's position. Do not fabricate case "
        "citations you are not reasonably confident exist; if uncertain, note the "
        "precedent as illustrative rather than binding."
    )

    user_prompt = (
        f"The client represents the '{client_role}' side.\n\n"
        f"PARTIES:\n{json.dumps(parties, indent=2)}\n\n"
        f"TIMELINE:\n{json.dumps(timeline, indent=2)}\n\n"
        f"CASE DOCUMENT EXCERPT:\n{raw_text[:40000]}\n\n"
        "Determine:\n"
        "1. The correct court_level (e.g. Small Claims, District/Lower Court, High "
        "Court, Labor Tribunal, Family Court, Commercial Court) and court_type.\n"
        "2. A short reasoning paragraph justifying that routing decision.\n"
        "3. An estimated court filing fee, expressed as a percentage/formula or flat "
        "fee, plus a list of filing_requirements (documents needed to file).\n"
        "4. A list of 2-5 relevant precedents with case_name, citation, relevance, "
        "and supports_client (true/false)."
    )

    try:
        llm = get_llm(temperature=0.1)
        structured_llm = llm.with_structured_output(_JurisdictionExtractionResult)
        result: _JurisdictionExtractionResult = structured_llm.invoke(
            [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]
        )
        jurisdiction = result.jurisdiction.model_dump()
        jurisdiction["estimated_court_fee"] = _reconcile_fee_estimate(
            jurisdiction.get("court_level", ""), jurisdiction.get("estimated_court_fee", "")
        )
        precedents = [p.model_dump() for p in result.precedents]
    except Exception as exc:  # noqa: BLE001
        errors.append(f"JurisdictionAndPrecedentNode: LLM extraction error - {exc}")
        jurisdiction = _empty_jurisdiction()
        precedents = []

    return {
        **state,
        "jurisdiction": jurisdiction,
        "precedents": precedents,
        "errors": errors,
        "status": "jurisdiction_routed",
    }


def _empty_jurisdiction() -> JurisdictionInfo:
    return {
        "court_level": "Unable to determine",
        "court_type": "Unable to determine",
        "reasoning": "Insufficient document text to determine jurisdiction.",
        "estimated_court_fee": "N/A",
        "filing_requirements": [],
    }


def _reconcile_fee_estimate(court_level: str, llm_fee_estimate: str) -> str:
    """Cross-checks the LLM's fee estimate against a static schedule for sanity."""
    key_map = {
        "small claims": "small_claims",
        "district": "district_civil",
        "lower court": "district_civil",
        "high court": "high_court_civil",
        "commercial": "commercial",
        "labor": "labor_tribunal",
        "labour": "labor_tribunal",
        "family": "family_court",
    }
    normalized = court_level.lower()
    for key, schedule_key in key_map.items():
        if key in normalized:
            schedule_fee = _COURT_FEE_SCHEDULE[schedule_key]
            if llm_fee_estimate and llm_fee_estimate.strip().lower() != "n/a":
                return f"{llm_fee_estimate} (reference schedule: {schedule_fee})"
            return schedule_fee
    return llm_fee_estimate or "Fee schedule not determined - consult local court registry"


# --------------------------------------------------------------------------
# NODE 4: LegalDraftingNode
# --------------------------------------------------------------------------

def LegalDraftingNode(state: CaseState) -> CaseState:
    """
    Writes a formal court speech/argument tailored to the identified court
    type, and produces a structured executive summary and party-positions
    analysis.
    """
    errors = state.get("errors", [])
    raw_text = state.get("raw_text", "")
    client_role = state.get("client_role", "Plaintiff / Victim")
    parties = state.get("parties", [])
    timeline = state.get("timeline", [])
    inconsistencies = state.get("inconsistencies", [])
    jurisdiction = state.get("jurisdiction", _empty_jurisdiction())
    precedents = state.get("precedents", [])

    if not raw_text:
        errors.append("LegalDraftingNode: no raw_text available to draft from")
        return {
            **state,
            "executive_summary": "",
            "party_positions": "",
            "formal_speech": "",
            "errors": errors,
        }

    system_prompt = (
        "You are a senior litigation counsel AI drafting materials for court "
        "submission. Your writing is formal, precise, persuasive but not inflammatory, "
        "and strictly grounded in the facts and precedents provided. You represent "
        f"the '{client_role}' side of this dispute and must advocate for that side "
        "while remaining professionally accurate about the facts, including known "
        "inconsistencies."
    )

    user_prompt = (
        f"COURT ROUTING: {jurisdiction.get('court_level')} - {jurisdiction.get('court_type')}\n"
        f"REASONING: {jurisdiction.get('reasoning')}\n\n"
        f"PARTIES:\n{json.dumps(parties, indent=2)}\n\n"
        f"TIMELINE:\n{json.dumps(timeline, indent=2)}\n\n"
        f"INCONSISTENCIES:\n{json.dumps(inconsistencies, indent=2)}\n\n"
        f"PRECEDENTS:\n{json.dumps(precedents, indent=2)}\n\n"
        "Produce three outputs:\n"
        "1. executive_summary: a concise (150-250 word) fact-sheet style summary of "
        "the case suitable for a case file cover page.\n"
        "2. party_positions: a structured comparison of each party's position and "
        "the core dispute, written in clearly separated paragraphs per party.\n"
        "3. formal_speech: a full formal court oral argument / speech (600-1000 "
        "words) appropriate for the identified court level and type, opening with "
        "an address to the court, presenting the facts and timeline, addressing the "
        "identified inconsistencies where relevant, citing the given precedents by "
        "name, and closing with a clear prayer for relief on behalf of the client."
    )

    try:
        llm = get_llm(temperature=0.3)
        structured_llm = llm.with_structured_output(_DraftingResult)
        result: _DraftingResult = structured_llm.invoke(
            [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]
        )
        executive_summary = result.executive_summary
        party_positions = result.party_positions
        formal_speech = result.formal_speech
    except Exception as exc:  # noqa: BLE001
        errors.append(f"LegalDraftingNode: LLM drafting error - {exc}")
        executive_summary, party_positions, formal_speech = "", "", ""

    return {
        **state,
        "executive_summary": executive_summary,
        "party_positions": party_positions,
        "formal_speech": formal_speech,
        "errors": errors,
        "status": "complete",
    }


# --------------------------------------------------------------------------
# GRAPH ASSEMBLY
# --------------------------------------------------------------------------

def build_case_graph():
    """Builds and compiles the LangGraph StateGraph for the case pipeline."""
    graph = StateGraph(CaseState)

    graph.add_node("document_parser", DocumentParserNode)
    graph.add_node("timeline_party_extractor", TimelineAndPartyExtractorNode)
    graph.add_node("jurisdiction_precedent", JurisdictionAndPrecedentNode)
    graph.add_node("legal_drafting", LegalDraftingNode)

    graph.set_entry_point("document_parser")
    graph.add_edge("document_parser", "timeline_party_extractor")
    graph.add_edge("timeline_party_extractor", "jurisdiction_precedent")
    graph.add_edge("jurisdiction_precedent", "legal_drafting")
    graph.add_edge("legal_drafting", END)

    return graph.compile()


def run_case_pipeline(file_path: str, file_name: str, client_role: str, file_id: str) -> CaseState:
    """Convenience entry point used by the FastAPI backend."""
    compiled_graph = build_case_graph()
    initial_state: CaseState = {
        "file_path": file_path,
        "file_name": file_name,
        "client_role": client_role,
        "file_id": file_id,
        "errors": [],
    }
    final_state = compiled_graph.invoke(initial_state)
    return final_state
